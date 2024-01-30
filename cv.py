from docx import Document

def create_cv(name, email, phone, education, experience, skills):
    # Create a new Document
    doc = Document()

    # Add a heading with the name
    doc.add_heading(name, level=1)

    # Add contact information
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"Phone: {phone}")

    # Add education section
    doc.add_heading('Education', level=2)
    for edu in education:
        doc.add_paragraph(f"{edu['degree']} in {edu['major']}")
        doc.add_paragraph(f"{edu['university']}, {edu['year']}")

    # Add experience section
    doc.add_heading('Experience', level=2)
    for exp in experience:
        doc.add_paragraph(f"{exp['position']}, {exp['company']}")
        doc.add_paragraph(f"{exp['start_date']} - {exp['end_date']}")
        doc.add_paragraph(exp['description'])

    # Add skills section
    doc.add_heading('Skills', level=2)
    doc.add_paragraph(', '.join(skills))

    # Save the document
    doc.save('my_cv.docx')

if __name__ == "__main__":
    # Example data
    my_name = "John Doe"
    my_email = "john.doe@example.com"
    my_phone = "+1234567890"
    
    my_education = [
        {'degree': 'Bachelor of Science', 'major': 'Computer Science', 'university': 'XYZ University', 'year': '2020'},
        # Add more education entries as needed
    ]

    my_experience = [
        {'position': 'Software Engineer', 'company': 'ABC Corp', 'start_date': '2020', 'end_date': '2022', 'description': 'Worked on various projects'},
        # Add more experience entries as needed
    ]

    my_skills = ['Python', 'Java', 'Web Development', 'Problem Solving']

    # Create the CV
    create_cv(my_name, my_email, my_phone, my_education, my_experience, my_skills)
